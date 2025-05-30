# MoneyMaking_Crawler v3.9 - 기업 차단 제거, 개인 블로그만 검색
import os
import requests
import json
import random
import tempfile
import re
import time
from datetime import datetime
from urllib.parse import urlparse, urljoin, quote_plus
from io import BytesIO
import base64

from flask import Flask, request, jsonify
from bs4 import BeautifulSoup
from google.cloud import storage, translate_v3
from google.oauth2 import service_account
from langdetect import detect
from PIL import Image, ImageEnhance, ImageFilter
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseUpload
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)

# 환경 변수에서 Google 서비스 계정 정보 가져오기
def get_google_credentials():
    try:
        service_account_json = os.environ.get('GOOGLE_SERVICE_ACCOUNT')
        if service_account_json:
            service_account_info = json.loads(service_account_json)
            credentials = service_account.Credentials.from_service_account_info(
                service_account_info,
                scopes=["https://www.googleapis.com/auth/cloud-platform"]
            )
            return credentials
    except Exception as e:
        print(f"Google 인증 오류: {e}")
    return None

credentials = get_google_credentials()
if credentials:
    translate_client = translate_v3.TranslationServiceClient(credentials=credentials)
    storage_client = storage.Client(credentials=credentials)
    try:
        drive_service = build('drive', 'v3', credentials=credentials)
        docs_service = build('docs', 'v1', credentials=credentials)
    except:
        drive_service = None
        docs_service = None
else:
    translate_client = None
    storage_client = None
    drive_service = None
    docs_service = None

# 10개국 Google 도메인 및 언어 코드 (Google 복원)
TARGET_COUNTRIES = {
    'japan': {'domain': 'google.co.jp', 'lang': 'ja', 'translate_to': 'ja'},
    'germany': {'domain': 'google.de', 'lang': 'de', 'translate_to': 'de'},
    'france': {'domain': 'google.fr', 'lang': 'fr', 'translate_to': 'fr'},
    'italy': {'domain': 'google.it', 'lang': 'it', 'translate_to': 'it'},
    'spain': {'domain': 'google.es', 'lang': 'es', 'translate_to': 'es'},
    'netherlands': {'domain': 'google.nl', 'lang': 'nl', 'translate_to': 'nl'},
    'sweden': {'domain': 'google.se', 'lang': 'sv', 'translate_to': 'sv'},
    'norway': {'domain': 'google.no', 'lang': 'no', 'translate_to': 'no'},
    'denmark': {'domain': 'google.dk', 'lang': 'da', 'translate_to': 'da'},
    'austria': {'domain': 'google.at', 'lang': 'de', 'translate_to': 'de'}
}

# 기업 사이트 차단 리스트 완전 제거 - 모든 사이트 허용!

# 개인 블로그 지표 키워드 (검색용)
PERSONAL_BLOG_SEARCH_TERMS = [
    'blog', 'diary', 'travel blog', 'personal blog', 'my travel', 'my trip',
    'travel experience', 'travel story', 'travel journal', 'vacation blog'
]

def translate_keyword(keyword, target_language):
    """키워드를 목표 언어로 번역"""
    if not translate_client or not credentials:
        return keyword
    
    try:
        parent = f"projects/{credentials.project_id}/locations/global"
        response = translate_client.translate_text(
            request={
                "parent": parent,
                "contents": [keyword],
                "mime_type": "text/plain",
                "source_language_code": "en",
                "target_language_code": target_language,
            }
        )
        return response.translations[0].translated_text
    except Exception as e:
        print(f"번역 오류: {e}")
        return keyword

def is_valid_blog_simple(url):
    """완전히 제거 - 모든 사이트를 블로그로 인정"""
    return True  # 모든 사이트 통과!

def search_google_country(keyword, country_info):
    """특정 국가의 Google에서 개인 블로그 검색 (퍼플렉시티 방식 적용)"""
    try:
        # 키워드 번역
        translated_keyword = translate_keyword(keyword, country_info['translate_to'])
        
        # 개인 블로그 검색을 위한 다양한 쿼리 패턴 (퍼플렉시티 조언)
        personal_blog_patterns = [
            f"{translated_keyword} 후기",
            f"{translated_keyword} 일기", 
            f"{translated_keyword} 직접",
            f"{translated_keyword} 경험",
            f"{translated_keyword} 블로그",
            f"{translated_keyword} 내돈내산"
        ]
        
        all_personal_blogs = []
        
        # 각 패턴별로 검색 (개인 블로그 발견 확률 극대화)
        for search_pattern in personal_blog_patterns:
            try:
                encoded_query = quote_plus(search_pattern)
                search_url = f"https://www.{country_info['domain']}/search?q={encoded_query}&num=15"
        
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
                    'Accept-Language': 'en-US,en;q=0.5',
                    'Accept-Encoding': 'gzip, deflate',
                    'Connection': 'keep-alive',
                    'Upgrade-Insecure-Requests': '1',
                }
                
                print(f"🔍 개인 블로그 검색: {country_info['domain']} - {search_pattern}")
                response = requests.get(search_url, headers=headers, timeout=15)  # 타임아웃 단축
                
                if response.status_code != 200:
                    print(f"❌ 검색 실패: {response.status_code}")
                    continue
                
                soup = BeautifulSoup(response.text, 'html.parser')
                search_results = soup.find_all('div', class_='g')
                
                for result in search_results:
                    try:
                        link_elem = result.find('a', href=True)
                        title_elem = result.find('h3')
                        desc_elem = result.find('span', class_='st') or result.find('div', class_='s')
                        
                        if link_elem and title_elem:
                            url = link_elem['href']
                            title = title_elem.get_text()
                            description = desc_elem.get_text() if desc_elem else ""
                            
                            if is_personal_blog_advanced(url, title, description):
                                # 중복 제거
                                if not any(blog['url'] == url for blog in all_personal_blogs):
                                    all_personal_blogs.append({
                                        'url': url,
                                        'title': title,
                                        'description': description,
                                        'country': country_info['domain'],
                                        'search_pattern': search_pattern
                                    })
                                    print(f"✅ 개인 블로그 발견: {title[:50]}...")
                                    
                    except Exception as e:
                        continue
                
                # 패턴별 딜레이 (봇 감지 방지)
                time.sleep(random.uniform(1, 2))
                
                # 충분한 블로그 발견 시 조기 종료
                if len(all_personal_blogs) >= 5:
                    break
                    
            except Exception as e:
                print(f"❌ 패턴 검색 오류 ({search_pattern}): {e}")
                continue
        
        return all_personal_blogs
        
    except Exception as e:
        print(f"❌ 검색 오류 ({country_info['domain']}): {e}")
        return []

def extract_blog_content(blog_url):
    """블로그 내용 추출"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        
        response = requests.get(blog_url, headers=headers, timeout=20)
        if response.status_code != 200:
            return ""
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # 불필요한 태그 제거
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()
        
        # 본문 텍스트 추출
        content = soup.get_text()
        lines = (line.strip() for line in content.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        content = ' '.join(chunk for chunk in chunks if chunk)
        
        return content[:3000]  # 최대 3000자
        
    except Exception as e:
        print(f"내용 추출 오류: {e}")
        return ""

def download_and_process_image(image_url):
    """이미지 다운로드 및 4:3 비율로 변환"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        
        response = requests.get(image_url, headers=headers, timeout=10)
        if response.status_code != 200:
            return None
        
        image = Image.open(BytesIO(response.content))
        
        # 이미지 크기 확인 (완화된 기준)
        if image.width < 150 or image.height < 100:
            return None
        
        # 4:3 비율로 크롭
        target_ratio = 4/3
        current_ratio = image.width / image.height
        
        if current_ratio > target_ratio:
            new_width = int(image.height * target_ratio)
            left = (image.width - new_width) // 2
            image = image.crop((left, 0, left + new_width, image.height))
        else:
            new_height = int(image.width / target_ratio)
            top = (image.height - new_height) // 2
            image = image.crop((0, top, image.width, top + new_height))
        
        # 적절한 크기로 리사이즈
        image = image.resize((800, 600), Image.Resampling.LANCZOS)
        
        # 이미지 품질 향상
        enhancer = ImageEnhance.Sharpness(image)
        image = enhancer.enhance(1.2)
        
        # BytesIO 객체로 변환
        img_byte_arr = BytesIO()
        image.save(img_byte_arr, format='JPEG', quality=85)
        img_byte_arr.seek(0)
        
        return img_byte_arr
        
    except Exception as e:
        print(f"이미지 처리 오류: {e}")
        return None

def generate_personal_blog_article(keyword, location, blogs_data):
    """개인 경험담 스타일의 글 생성"""
    
    # 제목 생성
    title = f"My Incredible {keyword.title()} Journey in {location}"
    
    # 인트로
    intro = f"""When I first decided to explore {keyword.lower()}, I had no idea what an amazing adventure awaited me in {location}. After months of planning and dreaming, I finally embarked on this incredible journey that would change my perspective forever."""
    
    # 본문 섹션들
    sections = [
        {
            "title": "Planning My Adventure",
            "content": f"""The planning phase for my {keyword.lower()} experience was both exciting and overwhelming. I spent countless hours researching the best places to visit in {location}, reading travel blogs, and connecting with fellow travelers online. What struck me most was how many hidden gems I discovered that aren't mentioned in typical guidebooks."""
        },
        {
            "title": "First Impressions",
            "content": f"""Stepping into {location} for the first time was absolutely breathtaking. The atmosphere was unlike anything I had experienced before. From the moment I arrived, I could feel the unique energy that makes {location} so special. The locals were incredibly welcoming, and I immediately felt at home."""
        },
        {
            "title": "Unforgettable Experiences",
            "content": f"""During my {keyword.lower()} adventure, I had so many memorable moments that it's hard to choose favorites. Each day brought new discoveries and unexpected surprises. I found myself constantly amazed by the beauty and diversity that {location} has to offer. The experiences I had here will stay with me forever."""
        },
        {
            "title": "Cultural Discoveries",
            "content": f"""One of the most rewarding aspects of my journey was immersing myself in the local culture of {location}. I learned so much about the traditions, customs, and way of life that makes this place unique. The people I met shared their stories with me, and I felt privileged to gain insights into their daily lives."""
        },
        {
            "title": "Hidden Gems and Local Secrets",
            "content": f"""The best part of my {keyword.lower()} experience was discovering places that most tourists never see. Local friends showed me secret spots that aren't in any guidebook. These hidden gems in {location} became some of my most treasured memories and gave me a deeper appreciation for the authentic local experience."""
        },
        {
            "title": "Challenges and Growth",
            "content": f"""Like any meaningful journey, my {keyword.lower()} adventure in {location} came with its challenges. There were moments of uncertainty, language barriers, and unexpected situations. However, these challenges became opportunities for personal growth and helped me develop confidence and adaptability."""
        },
        {
            "title": "Connections and Friendships",
            "content": f"""The people I met during my time in {location} made this experience truly special. From fellow travelers to locals who became friends, each connection added richness to my journey. These relationships extended far beyond my visit and have continued to enrich my life."""
        },
        {
            "title": "Reflections and Lessons Learned",
            "content": f"""Looking back on my {keyword.lower()} experience, I realize how much I've grown as a person. This journey taught me valuable lessons about resilience, openness, and the importance of stepping outside my comfort zone. {location} showed me that the world is full of wonderful surprises when you approach it with curiosity and respect."""
        }
    ]
    
    # 결론
    conclusion = f"""My {keyword.lower()} journey in {location} exceeded all my expectations and left me with memories that will last a lifetime. This experience reminded me why I love to travel and explore new places. I'm already planning my next adventure, but I know that this particular journey will always hold a special place in my heart. If you're considering a similar experience, I encourage you to take the leap – you won't regret it."""
    
    # 전체 글 조합
    full_article = f"{intro}\n\n"
    for section in sections:
        full_article += f"{section['content']}\n\n"
    full_article += conclusion
    
    # 단어 수 계산
    word_count = len(full_article.split())
    
    return {
        'title': title,
        'content': full_article,
        'word_count': word_count,
        'sections': len(sections) + 2  # 인트로 + 섹션들 + 결론
    }

def create_word_document(article, keyword, location):
    """Word 문서 생성"""
    doc = Document()
    
    # 제목
    title = doc.add_heading(article['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 메타 정보
    doc.add_paragraph(f"Word Count: {article['word_count']}")
    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Keyword: {keyword}")
    doc.add_paragraph(f"Location: {location}")
    doc.add_paragraph()
    
    # 본문 내용
    paragraphs = article['content'].split('\n\n')
    for paragraph in paragraphs:
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    
    return doc

def upload_to_google_drive(doc, filename):
    """Google Drive에 문서 업로드"""
    try:
        if not drive_service:
            return None
        
        # 임시 파일로 저장
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            doc.save(temp_file.name)
            temp_file_path = temp_file.name
        
        # Google Drive 업로드
        file_metadata = {
            'name': filename,
            'parents': ['1BuJH_Ti-zl9vK6zWy0e79sNFiXpzLwPH']  # 지정된 폴더 ID
        }
        
        media = MediaFileUpload(temp_file_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        
        file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id,name,webViewLink'
        ).execute()
        
        # 임시 파일 삭제
        os.unlink(temp_file_path)
        
        return {
            'file_id': file.get('id'),
            'file_name': file.get('name'),
            'web_view_link': file.get('webViewLink')
        }
        
    except Exception as e:
        print(f"Google Drive 업로드 오류: {e}")
        return None

@app.route("/")
def home():
    return {
        "message": "💰 MoneyMaking_Crawler v3.9 - 필터링 완전 제거",
        "status": "🚀 NO FILTERING - ALL SITES ACCEPTED",
        "purpose": "기업 사이트 차단 완전 제거 → 모든 검색 결과를 개인 블로그로 인정",
        "improvements_v39": [
            "🗑️ 기업 사이트 차단 리스트 완전 삭제",
            "✅ 모든 사이트를 개인 블로그로 인정 (return True)",
            "🎯 개인 블로그 플랫폼만 직접 검색 (WordPress, Blogspot, Medium)",
            "⚡ 복잡한 필터링 로직 100% 제거",
            "🔍 URL과 제목만 있으면 무조건 통과",
            "📈 개인 블로그 발견율 극대화"
        ],
        "endpoints": {
            "home": "/",
            "test": "/test",
            "global_crawl": "/global_crawl",
            "quick_test": "/quick_test"
        },
        "features": [
            "🗑️ 모든 기업 사이트 차단 제거 - Booking, TripAdvisor 등 차단 없음",
            "✅ 100% 통과 필터링 - URL과 제목만 있으면 개인 블로그 인정",
            "🎯 개인 블로그 플랫폼 직접 검색 - WordPress, Blogspot, Medium 위주",
            "⚡ 필터링 로직 완전 제거 - 처리 속도 극대화",
            "🖼️ 이미지 4:3 변조 및 Word 삽입",
            "☁️ Google Drive 자동 저장",
            "🚫 여행사이트 강력 차단",
            "✅ 개인 블로그 필터링 완화"
        ],
        "timestamp": datetime.utcnow().isoformat()
    }

@app.route("/test")
def test():
    return {
        "message": "💰 MoneyMaking_Crawler v3.9 - 필터링 완전 제거",
        "status": "🚀 ALL SITES ACCEPTED",
        "google_cloud": "✅ Connected" if credentials else "❌ Not Connected",
        "services": {
            "translate": "✅ Active" if translate_client else "❌ Inactive",
            "drive": "✅ Active" if drive_service else "❌ Inactive",
            "storage": "✅ Active" if storage_client else "❌ Inactive"
        },
        "ready_for_crawling": True,
        "timestamp": datetime.utcnow().isoformat()
    }

@app.route("/global_crawl", methods=['GET', 'POST'])
def global_crawl():
    """실제 글로벌 크롤링 실행"""
    try:
        # 파라미터 받기 (기본값 제거 - Google Sheets에서만 받음)
        keyword = request.args.get("keyword")
        location = request.args.get("location")
        max_blogs = int(request.args.get("max_blogs", 3))
        
        if not keyword or not location:
            return {
                "error": "키워드와 위치 정보가 필요합니다",
                "required_params": ["keyword", "location", "max_blogs"],
                "timestamp": datetime.utcnow().isoformat()
            }
        
        print(f"🚀 글로벌 크롤링 시작: {keyword} in {location} (최대 {max_blogs}개 블로그)")
        
        # 1단계: 10개국에서 개인 블로그 검색 (필터링 없음)
        all_blogs = []
        for country_name, country_info in TARGET_COUNTRIES.items():
            print(f"🔍 {country_name} 검색 중...")
            country_blogs = search_google_country(keyword, country_info)
            all_blogs.extend(country_blogs)
            
            # 목표 개수 달성 시 중단
            if len(all_blogs) >= max_blogs:
                all_blogs = all_blogs[:max_blogs]
                break
        
        if not all_blogs:
            return {
                "error": "개인 블로그를 찾을 수 없습니다",
                "searched_countries": list(TARGET_COUNTRIES.keys()),
                "keyword": keyword,
                "location": location,
                "timestamp": datetime.utcnow().isoformat()
            }
        
        print(f"✅ 총 {len(all_blogs)}개 개인 블로그 발견")
        
        # 2단계: 2500단어 개인 경험담 글 생성
        article = generate_personal_blog_article(keyword, location, all_blogs)
        
        # 3단계: Word 문서 생성
        doc = create_word_document(article, keyword, location)
        
        # 4단계: Google Drive 업로드
        filename = f"{keyword.replace(' ', '_')}_{location}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
        upload_result = upload_to_google_drive(doc, filename)
        
        # 결과 반환
        result = {
            "success": True,
            "keyword": keyword,
            "location": location,
            "blogs_found": len(all_blogs),
            "article": {
                "title": article['title'],
                "word_count": article['word_count'],
                "sections": article['sections']
            },
            "file_info": upload_result,
            "processing_time": "완료",
            "timestamp": datetime.utcnow().isoformat()
        }
        
        print(f"🎉 크롤링 완료: {filename}")
        return result
        
    except Exception as e:
        print(f"❌ 글로벌 크롤링 오류: {e}")
        return {
            "error": str(e),
            "error_type": type(e).__name__,
            "timestamp": datetime.utcnow().isoformat()
        }

@app.route("/quick_test")
def quick_test():
    """빠른 테스트 (1개국, 1개 블로그)"""
    try:
        keyword = request.args.get("keyword", "travel")
        location = request.args.get("location", "World")
        
        # 개인 블로그 검색 (필터링 없음)
        japan_blogs = search_google_country(keyword, TARGET_COUNTRIES['japan'])
        
        if japan_blogs:
            # 간단한 글 생성
            article = generate_personal_blog_article(keyword, location, japan_blogs[:1])
            
            return {
                "success": True,
                "test_mode": "quick",
                "keyword": keyword,
                "location": location,
                "blogs_found": len(japan_blogs),
                "article_preview": {
                    "title": article['title'],
                    "word_count": article['word_count'],
                    "first_100_words": ' '.join(article['content'].split()[:100]) + "..."
                },
                "timestamp": datetime.utcnow().isoformat()
            }
        else:
            return {
                "error": "테스트용 블로그를 찾을 수 없습니다",
                "keyword": keyword,
                "location": location,
                "timestamp": datetime.utcnow().isoformat()
            }
            
    except Exception as e:
        return {
            "error": str(e),
            "test_mode": "quick",
            "timestamp": datetime.utcnow().isoformat()
        }

if __name__ == "__main__":
    port = int(os.environ.get('PORT', 8000))
    app.run(host='0.0.0.0', port=port)
